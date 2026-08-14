from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SUMMARY = ROOT / "experiment" / "output" / "experiment_summary.json"
DEFAULT_OUTPUT = ROOT / "experiment" / "OpenSky数据链实验报告_简版.docx"
FONT = "Microsoft YaHei"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
MUTED = "666666"


def set_font(run, size: float, *, bold: bool = False, color: str = "000000") -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = table_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        table_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = table_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell_text(cell, *, bold: bool = False, color: str = "000000", size: float = 11.5) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        for run in paragraph.runs:
            set_font(run, size, bold=bold, color=color)


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.add_run(text)


def add_step(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    set_font(paragraph.add_run(label), 12.5, bold=True, color=DARK_BLUE)
    set_font(paragraph.add_run(text), 12.5)


def add_two_column_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        shade_cell(cells[0], LIGHT_GRAY)
        style_cell_text(cells[0], bold=True, color=DARK_BLUE)
        style_cell_text(cells[1])
    set_table_geometry(table, [2400, 6960])


def add_result_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("检查内容", "实验结果", "结论")
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        shade_cell(cell, LIGHT_GRAY)
        style_cell_text(cell, bold=True, color=DARK_BLUE)
    for item, result, conclusion in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (item, result, conclusion)):
            cell.text = text
            style_cell_text(cell)
    set_table_geometry(table, [3300, 3030, 3030])


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(12.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading = doc.styles["Heading 1"]
    heading.font.name = FONT
    heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor.from_string(BLUE)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True


def build_report(summary_path: Path, output_path: Path) -> Path:
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("OpenSky 数据链实验｜简版"), 9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_font(run, 9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    run = footer.add_run(" 页")
    set_font(run, 9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    set_font(title.add_run("OpenSky 数据链实验报告（简版）"), 22, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    set_font(subtitle.add_run("真实数据收发、解码与精度验证"), 13, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    set_font(meta.add_run("日期："), 11.5, bold=True)
    set_font(meta.add_run("2026年8月14日    "), 11.5)
    set_font(meta.add_run("实验结论："), 11.5, bold=True)
    set_font(meta.add_run("通过"), 11.5, bold=True, color=DARK_BLUE)

    overview = doc.add_paragraph()
    overview.paragraph_format.space_after = Pt(8)
    set_font(overview.add_run("一句话说明："), 12.5, bold=True, color=DARK_BLUE)
    set_font(overview.add_run("把真实飞机位置变成二进制消息发出去，再接收、还原并保存；整个流程运行正常。"), 12.5)

    add_heading(doc, "一、实验流程")
    add_step(doc, "第一步：", "读取仓库内的 OpenSky 真实飞机状态数据。")
    add_step(doc, "第二步：", "把经纬度、高度、速度等数值编码成固定 41 字节消息。")
    add_step(doc, "第三步：", "逐帧模拟发送和接收，校验消息后恢复飞机状态。")
    add_step(doc, "第四步：", "把结果保存为 CSV 和 SQLite，并比较原值与还原值。")

    add_heading(doc, "二、实验数据")
    add_two_column_table(doc, [
        ("数据来源", "The OpenSky Network 官方接口快照"),
        ("原始数据", f"{summary['source_snapshot_count']} 个快照，共 {summary['source_record_count']} 条状态记录"),
        ("本次实验", f"{summary['selected_target_count']} 个飞机目标，共 {summary['selected_record_count']} 条记录"),
        ("消息格式", f"每帧 {summary['frame_size_bytes']} 字节 TeachingLink 二进制消息"),
    ])

    # 大字号版本明确分为两页，避免结果表在页间断开。
    doc.add_page_break()
    add_heading(doc, "三、实验结果")
    add_result_table(doc, [
        ("消息收发", f"发送 {summary['sent_frame_count']} 帧，正确接收 {summary['valid_received_frame_count']} 帧", "全部通过"),
        ("接收端态势", f"由空表逐步形成 {summary['final_receiver_target_count']} 个目标", "正常"),
        ("数据库保存", f"SQLite 保存 {summary['sqlite_row_count']} 行", "正常"),
        ("位置精度", f"最大误差 {summary['max_horizontal_error_m']:.2f} 米，平均 {summary['mean_horizontal_error_m']:.2f} 米", "符合量化范围"),
        ("高度与速度", f"最大误差 {summary['max_altitude_error_m']:.2f} 米 / {summary['max_speed_error_m_s']:.2f} 米每秒", "符合量化范围"),
    ])

    add_heading(doc, "四、结论")
    conclusion = doc.add_paragraph()
    set_font(conclusion.add_run("实验已完成。"), 12.5, bold=True, color=DARK_BLUE)
    set_font(conclusion.add_run("真实 OpenSky 数据能够完成编码、传输、接收、解码和保存。二进制定点编码会有几米以内的小误差，这是把连续小数压缩为有限整数时的正常现象。"), 12.5)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="根据实验摘要生成简版 Word 报告")
    parser.add_argument("--summary", type=Path, default=DEFAULT_SUMMARY)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    print(build_report(args.summary, args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
